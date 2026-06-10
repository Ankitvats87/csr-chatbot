import csv
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.utils.env import Settings
from app.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class TextSegment:
    text: str
    page: int  # 1-indexed; for non-paginated docs, page == 1


def extract_segments(local_path: str, settings: Optional[Settings] = None) -> List[TextSegment]:
    """Returns one TextSegment per logical page (PDFs) or one segment for the whole file."""
    ext = Path(local_path).suffix.lower().lstrip(".")
    
    # Use LlamaParse if API key is configured for PDF/DOCX
    if settings and settings.llama_cloud_api_key and ext in ["pdf", "docx"]:
        return _extract_llamaparse(local_path, settings.llama_cloud_api_key)

    if ext == "pdf":
        return _extract_pdf(local_path)
    if ext == "docx":
        return _extract_docx(local_path)
    if ext == "txt":
        return _extract_txt(local_path)
    if ext == "csv":
        return _extract_csv(local_path)
    logger.warning("unsupported extension", extra={"ext": ext, "path": local_path})
    return []


def _extract_llamaparse(path: str, api_key: str) -> List[TextSegment]:
    try:
        from llama_parse import LlamaParse
    except ImportError:
        logger.error("llama-parse library is not installed, falling back to local parsers")
        ext = Path(path).suffix.lower().lstrip(".")
        if ext == "pdf":
            return _extract_pdf(path)
        elif ext == "docx":
            return _extract_docx(path)
        return []

    logger.info("Extracting structured markdown layout using LlamaParse...", extra={"path": path})
    try:
        parser = LlamaParse(
            api_key=api_key,
            result_type="markdown",
            verbose=False,
            language="en",
            split_by_page=True,
            system_prompt=(
                "This is a Corporate Social Responsibility (CSR) committee meeting agenda or minutes document. "
                "Transcribe ONLY the text that is actually present on each page. "
                "Extract tables as markdown tables, preserving the exact project names, NGO names, figures and headers "
                "that appear in the document. "
                "STRICT RULE: NEVER invent, fabricate, or insert example/sample/placeholder content of any kind "
                "(such as 'Project A', 'NGO A', 'John Doe', or round example figures). "
                "If a page is blank, unreadable, or contains only images you cannot read, return nothing for that page."
            ),
        )
        documents = parser.load_data(path)
        out: List[TextSegment] = []
        for i, doc in enumerate(documents, start=1):
            text = (doc.text or "").strip()
            if text:
                out.append(TextSegment(text=text, page=i))
        logger.info("LlamaParse extraction complete", extra={"path": path, "pages": len(out)})
        return out
    except Exception as e:
        logger.exception("LlamaParse extraction failed, falling back to local parsers", extra={"path": path, "err": str(e)})
        ext = Path(path).suffix.lower().lstrip(".")
        if ext == "pdf":
            return _extract_pdf(path)
        elif ext == "docx":
            return _extract_docx(path)
        return []


def _extract_pdf(path: str) -> List[TextSegment]:
    out: List[TextSegment] = []
    reader = PdfReader(path)
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            logger.warning("pdf page extract failed", extra={"path": path, "page": i, "err": str(e)})
            text = ""
        text = text.strip()
        if text:
            out.append(TextSegment(text=text, page=i))
    return out


def _extract_docx(path: str) -> List[TextSegment]:
    doc = DocxDocument(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                t = cell.text.strip()
                if t and (not cells_text or cells_text[-1] != t):
                    cells_text.append(t)
            if cells_text:
                parts.append(" | ".join(cells_text))
    text = "\n".join(parts).strip()
    return [TextSegment(text=text, page=1)] if text else []


def _extract_txt(path: str) -> List[TextSegment]:
    text = Path(path).read_text(encoding="utf-8", errors="ignore").strip()
    return [TextSegment(text=text, page=1)] if text else []


def _extract_csv(path: str) -> List[TextSegment]:
    rows: List[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(", ".join(c.strip() for c in row))
    text = "\n".join(rows).strip()
    return [TextSegment(text=text, page=1)] if text else []


class Chunker:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def chunk_file(self, local_path: str) -> List[Tuple[str, int]]:
        """Returns a list of (chunk_text, page) tuples."""
        segments = extract_segments(local_path, self.settings)
        out: List[Tuple[str, int]] = []
        for seg in segments:
            for piece in self.splitter.split_text(seg.text):
                piece = piece.strip()
                if piece:
                    out.append((piece, seg.page))
        return out
